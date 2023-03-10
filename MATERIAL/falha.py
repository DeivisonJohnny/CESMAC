from docx import Document as dc

nome = input("Qual o seu nome? ")

assunto = input(f"Olá {nome} sobre qual assunto você gostaria de saber,\nSistema Operacionais ou Arquitetura de dados?")

sistemaSO = dc("sistema.docx")
arquiteturaPC = dc("arquitetura.docx")

while assunto:
        if(assunto == 'sistemas operacionais' or assunto == 'so' or assunto == 'sistemas'):    
            for sistemas in sistemaSO.paragraphs:
                print('\nLeia um texto sobre Sistemas Operacionais abaixo :)\n\n',sistemas.text)
                assunto = input('\nMais algum assunto? ')
                if(assunto == 'sim'): 
                          assunto = input('O que mais gostaria de saber? ')
                          if(assunto == 'arquitetura de dados' or assunto == 'at' or assunto == 'arquitetura'):
                               print(arquitetura.text)
                               assunto = input('Gostaria de ver também sobre Sistemas Operacionais? ')
                          else:
                               break
                else: 
                     print('Fim do programa')
                     break

        else:
            if(assunto == 'arquitetura de dados'):
                 for arquitetura in arquiteturaPC.paragraphs:
                    print('Leia um texto sobre sobre arquitetura de computadores abaixo :)\n\n',arquitetura.text)
                    assunto = input('Mais algum assunto? ')
                    if(assunto == 'não' or 'nao' or 'exit' or 'fechar'):
                        break
                    

            else: 
                 
               if(assunto == 'sistemas operacionais'):
                    for sistemas in sistemaSO.paragraphs:
                         print('\n',sistemas.text)

               else:
                    if(assunto == 'arquitetura de computadores' or 'arquitetura'):
                         for arquitetura in arquiteturaPC.paragraphs:
                              print(arquitetura.text)
                         else: 
                              break

               print('Desculpe-me, não entendi o seu comando.\n')

               assunto = input("\nMe diga novamento sobre qual assunto gostaria de saber.\n Entre SISTEMAS OPERACIONAIS E ARQUITETURA DE COMPUTADORES? ")

                 
            
            fim = input('Mais alguma coisa? ')
            if(fim == 'sim'):
                 break
            else:
                 assunto = input('Qual assunto? ')
                 

        